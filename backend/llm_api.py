import os
import json
import re
import time
import sys
from datetime import datetime
from tempfile import NamedTemporaryFile
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Dict, Any
import uvicorn

# --- KÜTÜPHANELER ---
try:
    import docx 
    import pytesseract
    from PIL import Image
    import pdfplumber
except ImportError:
    pass

from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_openai import ChatOpenAI, OpenAIEmbeddings 
from langchain_core.prompts import PromptTemplate 
from langchain_core.output_parsers import StrOutputParser
from langchain_community.vectorstores import FAISS 
from langchain_core.documents import Document 

pytesseract.pytesseract.tesseract_cmd = '/opt/homebrew/bin/tesseract'

# --- API AYARLARI ---
class MatchResult(BaseModel):
    job_title: str
    general_score: float
    skill_match: float
    experience_match: float
    report_summary: str

app = FastAPI(title="Hibrit LLM CV Eşleştirme API'si (Final)")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], 
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ====================================================================
# 1. VERİ YÜKLEME VE VEKTÖR İNDEKSİ (FAISS)
# ====================================================================

embedding_model = OpenAIEmbeddings(model="text-embedding-3-small")
GLOBAL_RETRIEVER = None
ALL_JOB_DATA = []

def load_data(file_name):
    try:
        with open(file_name, 'r', encoding='utf-8') as f:
            return json.load(f)
    except FileNotFoundError:
        print(f"KRİTİK HATA: {file_name} dosyası bulunamadı.")
        return [] 

def initialize_vector_store():
    global GLOBAL_RETRIEVER, ALL_JOB_DATA
    ALL_JOB_DATA = load_data('parsed_jobs_FINAL.json')
    if not ALL_JOB_DATA: return

    job_docs = []
    for job in ALL_JOB_DATA:
        weighted_content = (
            f"Pozisyon: {job['job_title']} " * 3 +
            f"Sektör: {job['sector']} " * 2 +
            f"Detaylar: {job.get('description', '')} {job.get('qualifications_raw', '')}"
        )
        job_docs.append(Document(page_content=weighted_content, metadata={"job_title": job['job_title']}))

    if job_docs:
        vector_store = FAISS.from_documents(job_docs, embedding_model)
        GLOBAL_RETRIEVER = vector_store.as_retriever(search_kwargs={"k": 100}) 
        print(f"INFO: FAISS Vektör İndeksi {len(job_docs)} ilan ile hazırlandı.")

initialize_vector_store()

# ====================================================================
# 2. HAFTA 4: PARSING, TEMİZLEME VE NORMALİZASYON
# ====================================================================

def clean_text(text: str) -> str:
    """Metin temizleme: Boşluklar ve anlamsız karakterler."""
    if not text: return ""
    text = re.sub(r'[\r\n\t]+', ' ', text) # Satır sonlarını boşluğa çevir
    text = re.sub(r'\s+', ' ', text)       # Çift boşlukları tek boşluğa indir
    return text.strip()

def normalize_titles(text: str) -> str:
    """Bölüm ve unvan normalizasyonu."""
    replacements = {
        r'end\.?\s*müh\.?': 'Endüstri Mühendisliği',
        r'bil\.?\s*müh\.?': 'Bilgisayar Mühendisliği',
        r'mak\.?\s*müh\.?': 'Makine Mühendisliği',
        r'yazılım uzm\.?': 'Yazılım Uzmanı',
        r'ik': 'İnsan Kaynakları'
    }
    text_lower = text.lower()
    for pattern, replacement in replacements.items():
        if re.search(pattern, text_lower):
            return replacement # Tam değişim yapıyoruz ki net olsun
    return text

def calculate_experience_duration(text: str) -> float:
    """Metindeki tarih aralıklarını bulup toplam deneyim yılını hesaplar."""
    year_pattern = re.findall(r'(\d{4})\s*-\s*(\d{4}|Devam|Halen|Present)', text)
    total_years = 0.0
    current_year = datetime.now().year
    
    for start, end in year_pattern:
        try:
            start_year = int(start)
            end_year = current_year if end in ['Devam', 'Halen', 'Present'] else int(end)
            diff = end_year - start_year
            if diff >= 0: total_years += diff
        except: continue
            
    if total_years == 0:
        text_year_match = re.search(r'(\d+)\s*(?:yıl|sene|year)', text.lower())
        if text_year_match: total_years = float(text_year_match.group(1))

    return total_years

def validate_is_real_cv(text_content: str) -> bool:
    """Dosyanın CV olup olmadığını kontrol eder (Kapı Bekçisi)."""
    if len(text_content) < 50: return False
    
    keywords = ["eğitim", "deneyim", "iş", "yetenekler", "beceriler", "education", "experience", "skills", "özgeçmiş", "cv", "iletişim", "university", "lise", "lisans"]
    found = [k for k in keywords if k in text_content.lower()]
    
    # En az 3 anahtar kelime geçmeli
    if len(found) >= 3: return True
    
    # Emin olamazsak LLM'e soralım
    try:
        llm = ChatOpenAI(model="gpt-3.5-turbo", temperature=0.0)
        prompt = f"Bu metin bir CV (Özgeçmiş) mi? Sadece 'EVET' veya 'HAYIR' yaz.\n\nMetin: {text_content[:500]}"
        res = llm.invoke(prompt).content.strip().upper()
        return "EVET" in res
    except: return True

def parse_cv_content(file_path):
    """Dosyadan metin okur (Gelişmiş OCR Desteği)."""
    pytesseract.pytesseract.tesseract_cmd = '/opt/homebrew/bin/tesseract'
    ext = os.path.splitext(file_path)[1].lower()
    text = ""
    
    try:
        if ext == '.pdf':
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    if page.extract_text(): text += page.extract_text() + "\n"
                    
        elif ext == '.docx':
            doc = docx.Document(file_path)
            text = "\n".join([p.text for p in doc.paragraphs])
            
        elif ext in ['.png', '.jpg', '.jpeg']:
            image = Image.open(file_path)
            
            # --- GELİŞMİŞ OCR AYARI ---
            # --psm 6: Sayfayı tek bir metin bloğu olarak gör (Sütunları karıştırmayı engeller)
            # lang='tur+eng': Türkçe ve İngilizce karakterleri tanı
            try:
                # Önce Tesseract'ın yolunu kontrol etmeye gerek yok, sistem yolunda olmalı.
                text = pytesseract.image_to_string(image, lang='tur+eng', config='--psm 6')
            except Exception as ocr_error:
                print(f"OCR Hatası (Tesseract Yüklü mü?): {ocr_error}")
                # Yedek: Sadece İngilizce dene (Bazen Türkçe paketi olmayabilir)
                text = pytesseract.image_to_string(image, lang='eng', config='--psm 6')

            print(f"INFO: OCR tamamlandı. Okunan karakter sayısı: {len(text)}")
            
        else: 
            with open(file_path, 'r', encoding='utf-8') as f: return f.read()
            
    except Exception as e:
        print(f"Genel Parse Hatası: {e}")
        return ""
    
    # Temizlik
    return re.sub(r'\s+', ' ', text).strip()

# ====================================================================
# 3. SKORLAMA MANTIĞI (AĞIRLIKLAR + CEZA)
# ====================================================================

def map_education_level(level_str):
    level_str = str(level_str).lower()
    if "doktora" in level_str: return 5
    if "yüksek lisans" in level_str: return 4
    if "üniversite" in level_str or "lisans" in level_str or "bachelor" in level_str or "mühendis" in level_str: return 3
    if "ön lisans" in level_str or "myo" in level_str: return 2
    if "lise" in level_str: return 1
    return 0

def calculate_rule_scores(query_cv, job_details):
    cv_lvl = map_education_level(query_cv.get('education_raw', ''))
    job_lvl = map_education_level(job_details.get('education_level', ''))
    s_edu = 1.0 if cv_lvl >= job_lvl else (0.5 if cv_lvl == job_lvl - 1 else 0.0)

    s_loc = 0.5 
    cv_loc = query_cv.get('location', '').lower()
    job_loc = job_details.get('location', '').lower()
    if cv_loc and job_loc:
        if cv_loc.split(',')[0] in job_loc: s_loc = 1.0
        elif "istanbul" in cv_loc and "istanbul" in job_loc: s_loc = 1.0

    s_sal = 1.0 if "uzman" in job_details['job_title'].lower() or "mühendis" in job_details['job_title'].lower() else 0.5
    
    return s_edu, s_loc, s_sal

def is_field_mismatch(cv_edu, job_title):
    """Mühendisi garson yapmayı engelleyen PYTHON KURALI."""
    cv_edu = cv_edu.lower()
    job_title = job_title.lower()
    
    if "mühendis" in cv_edu:
        allowed = ["mühendis", "yazılım", "software", "developer", "geliştirici", "uzman", "analist", "teknik", "ar-ge", "bilim"]
        if any(k in job_title for k in allowed): return False 
        
        banned = ["garson", "temizlik", "bellboy", "steward", "şoför", "kasiyer", "satış danışmanı", "resepsiyonist"]
        if any(k in job_title for k in banned): return True 

    return False

# ====================================================================
# 4. LLM RERANKING
# ====================================================================

def rank_with_llm_logic(query_cv, job_list):
    llm = ChatOpenAI(model="gpt-3.5-turbo", temperature=0.0) 
    
    # 👇 DÜZELTME: Prompt'a "Mühendislik = Yazılım" kuralını sert bir şekilde ekledik.
    prompt = PromptTemplate.from_template("""
    Sen Teknik Bir İşe Alım Uzmanısın. Aşağıdaki adayı iş ilanıyla karşılaştır.
    
    ADAY BÖLÜMÜ: {cv_edu}
    ADAY YETKİNLİKLERİ: {cv_text}
    
    İLAN BAŞLIĞI: {job_title}
    İLAN DETAYI: {job_desc}
    
    --- KRİTİK EŞLEŞTİRME KURALLARI ---
    1. **MÜHENDİSLİK KURALI:** Eğer aday "Bilgisayar Mühendisliği", "Yazılım Mühendisliği" veya "Bilişim" mezunuysa; "Yazılım Uzmanı", "Developer", "Geliştirici", "Tester", "Analist", "Engineer" ilanları ile **%100 AYNI ALANDADIR.** Asla alan uyuşmazlığı deme!
    
    2. **PUANLAMA (0-10):**
       - Alan/Bölüm tutuyorsa (Yukarıdaki kural): **En az 7 Puan** ver.
       - Alan tutuyor ama tecrübe eksikse (Junior vs Senior): **6 Puan** ver.
       - Alan tamamen alakasızsa (Örn: Mühendis -> Garson): **0 Puan** ver.

    3. **RAPORLAMA:** - Eğer puanı kırdıysan sebebini "Tecrübe eksikliği" veya "Teknik yetkinlik eksikliği" olarak belirt. "Alan uyuşmazlığı" deme.
    
    JSON ÇIKTISI VER: 
    {{ "uyum_skoru": [PUAN], "analiz_ozeti": "[KISA VE MANTIKLI AÇIKLAMA]" }}
    """)
    
    reranked = []
    chain = prompt | llm | StrOutputParser()
    
    # İlk 15 ilanı analiz et
    for job in job_list[:15]: 
        
        cv_text_short = (query_cv.get('experience_raw', '') + " " + query_cv.get('skills_raw', ''))[:500]
        
        try:
            response = chain.invoke({
                "cv_edu": query_cv.get('education_raw', 'Belirtilmemiş'),
                "cv_text": cv_text_short,
                "job_title": job['job_title'],
                "job_desc": job.get('description', '')[:300]
            })
            
            match = re.search(r'\{.*\}', response, re.DOTALL)
            if match:
                data = json.loads(match.group(0))
                llm_score = float(data.get('uyum_skoru', 0))
                reason = data.get('analiz_ozeti', 'Analiz edildi.')
            else:
                llm_score = 0; reason = "LLM Yanıt Hatası"
                
        except Exception as e:
            # print(f"LLM Hatası ({job['job_title']}): {e}") 
            llm_score = 0; reason = "API Hatası"

        # Python tarafında ekstra güvenlik (Mühendis -> Yazılım Uzmanı eşleşmesini garantiye al)
        if "bilgisayar" in query_cv.get('education_raw', '').lower() and "yazılım" in job['job_title'].lower():
            if llm_score < 5: # Eğer LLM hata yapıp düşük verdiyse düzelt
                llm_score = 7.0
                reason = "Bölüm ve pozisyon teknik olarak uyumlu (Otomatik Düzeltme)."

        # Alan Uyuşmazlığı Kontrolü (Sadece gerçekten alakasızlar için)
        if is_field_mismatch(query_cv.get('education_raw', ''), job['job_title']):
            llm_score = 0.0
            reason = f"ALAN UYUŞMAZLIĞI: {query_cv.get('education_raw')} -> {job['job_title']}"

        # Final Skor Hesapla
        s_edu, s_loc, s_sal = calculate_rule_scores(query_cv, job)
        final_score = (0.6 * (llm_score/10)) + (0.1 * s_edu) + (0.1 * s_loc) + (0.2 * s_sal)
        
        reranked.append({
            "job_title": job['job_title'],
            "general_score": final_score,
            "skill_match": llm_score / 10, 
            "experience_match": s_edu, 
            "report_summary": reason
        })
        
    reranked.sort(key=lambda x: x['general_score'], reverse=True)
    return reranked[:5]

# ====================================================================
# 5. API ENDPOINT
# ====================================================================

def quick_llm_parse(text):
    """CV'den Bilgi Çekme"""
    llm = ChatOpenAI(model="gpt-3.5-turbo", temperature=0.0)
    prompt = """Aşağıdaki CV metnini analiz et ve JSON formatında bilgileri çıkar:
    { "name": "Ad Soyad", "education_raw": "Bölüm/Eğitim (Örn: Bilgisayar Mühendisliği)", "location": "Şehir" }
    Metin: """ + text[:1500]
    
    try:
        res = llm.invoke(prompt).content
        match = re.search(r'\{.*\}', res, re.DOTALL)
        if match: return json.loads(match.group(0))
    except: pass
    return {"name": "Aday", "education_raw": "Belirtilmemiş", "location": "İstanbul"}

@app.post("/api/match_cv", response_model=List[MatchResult])
async def match_cv(file: UploadFile = File(...)):
    with NamedTemporaryFile(delete=False, suffix=f".{file.filename.split('.')[-1]}") as temp:
        temp.write(await file.read())
        temp_path = temp.name

    # 1. Parsing & Temizleme
    cv_text = parse_cv_content(temp_path)
    os.remove(temp_path)
    
    cv_text = clean_text(cv_text) # Temizleme fonksiyonu
    
    if not cv_text.strip():
        raise HTTPException(status_code=400, detail="Dosya boş veya okunamadı.")

    # 2. Validasyon (CV mi?)
    if not validate_is_real_cv(cv_text):
        raise HTTPException(status_code=400, detail="Bu dosya bir CV'ye benzemiyor.")

    # 3. Bilgi Çıkarımı ve Normalizasyon
    cv_data = quick_llm_parse(cv_text)
    
    # Normalizasyonları Uygula
    cv_data['education_raw'] = normalize_titles(cv_data.get('education_raw', ''))
    cv_data['experience_years'] = calculate_experience_duration(cv_text)
    
    cv_data['text'] = cv_text 
    cv_data['experience_raw'] = cv_text 
    cv_data['skills_raw'] = cv_text

    print(f"\n---> İŞLENEN CV: {cv_data.get('education_raw')} | Süre: {cv_data.get('experience_years')} Yıl")

    # 4. Arama ve Sıralama (Keyword Rescue + FAISS)
    candidate_jobs = []
    
    # Keyword Rescue (Mühendis Koruması)
    keywords = ["yazılım", "software", "developer", "mühendis", "bilgisayar", "bilişim"]
    cv_lower = cv_text.lower()
    
    if any(k in cv_lower for k in keywords):
        for job in ALL_JOB_DATA:
            job_str = (job['job_title'] + " " + job['sector']).lower()
            if any(k in job_str for k in keywords):
                if job not in candidate_jobs: candidate_jobs.append(job)

    # FAISS (Ek olarak)
    if GLOBAL_RETRIEVER:
        query = f"{cv_data['education_raw']} {cv_text[:500]}"
        relevant_docs = GLOBAL_RETRIEVER.invoke(query) 
        relevant_titles = [doc.metadata['job_title'] for doc in relevant_docs]
        for title in relevant_titles:
            job = next((j for j in ALL_JOB_DATA if j['job_title'] == title), None)
            if job and job not in candidate_jobs:
                candidate_jobs.append(job)

    if len(candidate_jobs) < 5: candidate_jobs = ALL_JOB_DATA

    final_results = rank_with_llm_logic(cv_data, candidate_jobs)
    
    return final_results

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)